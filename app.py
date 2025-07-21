# app.py (完全版)

# === 必要なライブラリをまとめてインポート ===
import streamlit as st
import os
import datetime
import io

# Google Cloud & Vertex AI 関連
from google.cloud import storage
import vertexai
from vertexai.generative_models import GenerativeModel, Part
from google.oauth2 import service_account

# Wordファイル生成関連
from docx import Document

# === アプリのUI（ユーザーインターフェース）部分 ===

st.set_page_config(page_title="AI議事録作成アプリ", layout="wide")

st.title("AI議事録作成アプリ 📄✍️")
st.markdown("""
音声ファイル（MP3, WAV, M4A, MP4）をアップロードするだけで、AIが自動で文字起こしを行い、議事録の要約・決定事項・ToDoリストを作成します。
""")
st.info("※処理には数分かかることがあります。ファイルアップロード後、しばらくお待ちください。")


# ファイルアップロードのウィジェット
uploaded_file = st.file_uploader(
    "議事録を作成したい音声・動画ファイルを選択してください。",
    type=["mp3", "wav", "m4a", "mp4"]
)

# === メインの処理は、すべてこの if ブロックの中に入れる ===
# ファイルがアップロードされたら、以下の処理を順番に実行する
if uploaded_file is not None:
    
    st.success(f"ファイル「{uploaded_file.name}」がアップロードされました。処理を開始します。")

    # --- 1. 認証と初期化 ---
    # `with st.spinner` を使うと、処理中にローディング表示を出せる
    with st.spinner("ステップ1/5: クラウド環境に接続しています..."):
        try:
            # Streamlit CloudのSecretsから認証情報を読み込む
            creds_dict = st.secrets["gcp_service_account"]
            creds = service_account.Credentials.from_service_account_info(creds_dict)
            storage_client = storage.Client(credentials=creds)
            
            # あなたの新しいプロジェクトIDと、AIを動かすリージョンを指定
            project_id = "final-minutes-app" # 👈 【重要】あなたの新しいGCPプロジェクトIDに書き換えてください
            location = "asia-northeast1"           # 東京リージョン
            
            vertexai.init(project=project_id, location=location, credentials=creds)

        except (FileNotFoundError, KeyError):
            # ローカル環境で実行する場合のフォールバック
            st.info("ローカル環境として実行します。（Secretsが見つかりませんでした）")
            storage_client = storage.Client()
            
            project_id = "final-minutes-app" # 👈 【重要】あなたの新しいGCPプロジェクトIDに書き換えてください
            location = "asia-northeast1"           # 東京リージョン
            
            vertexai.init(project=project_id, location=location)

    # --- 2. Google Cloud Storageへのファイルアップロード ---
    gcs_uri = None
    with st.spinner("ステップ2/5: 音声ファイルを安全なクラウドにアップロードしています..."):
        # あなたの新しいGCSバケット名を指定
        bucket_name = "gijiroku-final-bucket" # 👈 【重要】あなたの新しいGCSバケット名に書き換えてください
        
        bucket = storage_client.bucket(bucket_name)
        timestamp = datetime.datetime.now().strftime("%Y%m%d-%H%M%S")
        blob_name = f"{timestamp}-{uploaded_file.name}"
        
        blob = bucket.blob(blob_name)
        blob.upload_from_file(uploaded_file, content_type=uploaded_file.type)
        
        gcs_uri = f"gs://{bucket_name}/{blob_name}"
        st.write(f"✅ ファイルのアップロードが完了しました: `{gcs_uri}`")


    # --- 3. AIによる文字起こし ---
    transcribed_text = None
    with st.spinner("ステップ3/5: AIが音声を文字に変換しています...（この処理は数分かかることがあります）"):
       
     # 使用するAIモデルをシンプルに定義する
     model = GenerativeModel("gemini-1.5-pro-preview-0514") # 動作実績のあるProモデル

     # GCS上の音声ファイルを Part オブジェクトとして準備
     audio_file_part = Part.from_uri(mime_type=uploaded_file.type, uri=gcs_uri)

     # テキストプロンプトも、Part オブジェクトとして明示的に準備
     prompt_text = """
        あなたは会議の内容を文字起こしするプロフェッショナルです。
この音声ファイルの全ての音声を以下の情報と条件にしたがって、文字起こしして。

【事前情報】
1. 会議を行う会社の概要
・会社名：SCN（読み：エスシーエヌ）
「営業チーム」
・中尾（読み：なかお）（部長）
・岡田（読み：おかだ）（課長）
・松村（読み：まつむら）
・植田（読み：うえだ）
「技術チーム」
・白間（読み：はくま）（センター長）
・二見（読み：ふたみ）（係長）
・笠（読み：りゅう）
・佐藤（読み：さとう）
「運行チーム」
・野口（読み：のぐち）（係長）
・杉本（読み：すぎもと）
・井上（読み：いのうえ）
・長谷川（読み：はせがわ）

【 用語集】
以下の単語が頻繁に登場します。表記は必ずこれに統一してください。
1.システム名
・Force（読み：フォース）
・Force-CMS(読み：フォースシーエムエス)
・Force-player（読み：フォースプレイヤー）
・Force-Guard（読み：フォースガード）
・cosmo5000(読み：コスモごせん)
・MOJIC(読み：モジック)
・Aves(読み：エイビス)
・新Aves（読み：しんえいびす）
・Ibis（読み：あいびす）
・管理Player(読み：管理プレイヤー)

2,局名
・ACCS（読み：エーシーシーエス）
・ACTV（読み：エーシーティーブイ）
・CC9（読み：シーシーナイン）
・CCNC（読み：シーシーエヌシー）
・CCO（読み：シーシーオー）
・CCS（読み：シーシーエス）
・CNA（読み：シーエヌエー）
・CNW（読み：シーエヌダブリュー）
・CTT（読み：シーティーティー）
・HCN（読み：エイチシーエヌ）
・KCT（読み：ケーシーティー）
・KDT（読み：ケーディーティー）
・KYT（読み：ケーワイティー）
・LCV（読み：エルシーブイ）
・MCN（読み：エムシーエヌ）
・NCN（読み：エヌシーエヌ）
・OCTV（読み：オーシーティーブイ）
・TCC（読み：ティーシーシー）
・YCV（読み：ワイシーブイ）
・YOUテレビ（読み：ユーテレビ）
・ZTV（読み：ゼットティーブイ）
・cable-one（読み：ケーブルワン）
・chupicom（読み：チュピコム）
・waiwai（読み：ワイワイ）
・こしの都ケーブルテレビ（読み：こしのみやこケーブルテレビ）
・オプテージ（読み：オプテージ）
・ケーブルテレビ品川（読み：ケーブルテレビしながわ）
・ケーブルテレビ徳島（読み：ケーブルテレビとくしま）
・仙台CATV（読み：せんだいシーエーティーブイ）
・佐伯市（読み：さえきし）
・南部町（読み：なんぶちょう）
・多摩ケーブルネットワーク（読み：たまケーブルネットワーク）
・多摩テレビ（読み：たまテレビ）
・大山町（読み：だいせんちょう）
・日野町（読み：ひのちょう）
・知多メディアス（読み：ちたメディアス）
・神津島（読み：こうづしま）
・美祢市（読み：みねし）
・茅野市（読み：ちのし）
・長門市（読み：ながとし）
・J:com（読み：ジェイコム）
・伯耆町（読み：ほうきちょう）
・高知（読み：こうち）
・中海テレビ（読み：ちゅうかいテレビ）
・岡山（読み：おかやま）
・倉敷（くらしき）

3.専門用語
・マルチ画面（読み：まるちがめん）
・マルチビューワ
・eBase(読み；いーべーす)
・メディアキャスト
・線状降水帯（読み：せんじょうこうすいたい）
・LM90（読み：えるえむきゅうじゅう）
・セルテック
・JWA（読み：ジェイダブリューエー）
・お天気＆LIVEカメラ（読み：おてんきあんどらいぶかめら）
・サーバー移行（読み：さーばーいこう）
・道録装置（読み：どうろくそうち）
・メディアエッジ
・南部SANチャンネル（読み：なんぶさんチャンネル）
・レイアウト
・河川（読み：かせん）
・端末（読み：たんまつ）
・HP（読み：エイチピー）
・３V（読み：すりーぶい）
・１０４（読み：いちまるよん）

【条件】
・フィラーを除去する。
・【用語】の単語が頻繁に登場します。表記は必ず【用語】の表記に統一してください。
・文字起こしした文が対応する音声ファイルの時間を正確に記入する
・話者はspeaker1などと表記し、人名は特定しない。
・「speaker1, 00:00:47,よろしくお願いいたします。」のように表記を統一する

        """
        prompt_part = Part.from_text(prompt_text)
        
        # AIにリクエストを送信 (Partオブジェクトだけで構成されたリストを渡す)
        response = model.generate_content([audio_file_part, prompt_part])
        
        
        transcribed_text = response.text
        st.subheader("文字起こし結果")
        with st.expander("全文を表示する"):
            st.write(transcribed_text)


    # --- 4. AIによる議事録生成 ---
    generated_minutes = None
    with st.spinner("ステップ4/5: AIが文字起こし内容を要約し、議事録を作成しています..."):
        # 議事録生成用の、より詳細なプロンプトを作成
        prompt_for_minutes = f"""
        あなたは非常に優秀なアシスタントです。
        以下の会議の文字起こしテキストを元に、要点を的確に捉えたプロフェッショナルな議事録を作成してください。

        以下のフォーマットを厳密に守って、日本語で出力してください。

        # 議事録

        ## 1. この会議の要約
        （会議全体の目的と結論を3〜5行で簡潔に記述してください）

        ## 2. 決定事項
        （会議で具体的に決定された事項を、箇条書きでリストアップしてください。決定事項がない場合は「特になし」と記載してください。）
        - 決定事項A
        - 決定事項B

        ## 3. ToDoリスト（担当者と期限）
        （会議で発生したタスクを、[ ]のチェックボックス形式でリストアップしてください。誰が(担当者)、いつまでに行うか(期限)を必ず明記してください。タスクがない場合は「特になし」と記載してください。）
        - [ ] 〇〇の調査（担当：佐藤さん、期限：YYYY-MM-DD）
        - [ ] △△の資料作成（担当：鈴木さん、期限：YYYY-MM-DD）
        
        ---
        
        # 元の文字起こしテキスト
        
        {transcribed_text}
        """
        
        # 再びAIにリクエストを送信
        response_minutes = model.generate_content(prompt_for_minutes)
        generated_minutes = response_minutes.text
        
        st.subheader("生成された議事録")
        st.markdown(generated_minutes)


    # --- 5. Wordファイル出力 ---
    with st.spinner("ステップ5/5: Wordファイルを生成しています..."):
        document = Document()
        document.add_heading('AI自動生成 議事録', 0)
        
        # 生成された議事録をドキュメントに追加
        # st.markdownと同じように表示するために、簡単なパーサーを実装
        for line in generated_minutes.split('\n'):
            line = line.strip()
            if line.startswith('### '):
                document.add_heading(line.replace('### ', ''), level=3)
            elif line.startswith('## '):
                document.add_heading(line.replace('## ', ''), level=2)
            elif line.startswith('# '):
                document.add_heading(line.replace('# ', ''), level=1)
            elif line.startswith('- '):
                # 箇条書きのインデントも考慮
                p = document.add_paragraph(style='List Bullet')
                p.add_run(line.replace('- ', ''))
            elif line: # 空行は無視
                document.add_paragraph(line)

        document.add_page_break()
        document.add_heading('文字起こし全文', level=1)
        document.add_paragraph(transcribed_text)
        
        # Wordファイルをメモリ上にバイナリデータとして保存
        file_stream = io.BytesIO()
        document.save(file_stream)
        file_stream.seek(0) # ストリームの先頭にポインタを戻す
    
    st.success("🎉 すべての処理が完了しました！以下から議事録をダウンロードできます。")

    # ダウンロードボタンを表示
    st.download_button(
        label="議事録をWordファイルでダウンロード",
        data=file_stream,
        file_name=f"議事録_{os.path.splitext(uploaded_file.name)[0]}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )